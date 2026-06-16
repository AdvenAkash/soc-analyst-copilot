#!/usr/bin/env python3
"""
Generate a Word document (and optionally PDF) from a directory of code files.

Usage:
    python3 scripts/code_to_docx.py <directory> [--team TEAM_NAME] [--output FILE.docx] [--pdf]

Examples:
    # Generate .docx
    python3 scripts/code_to_docx.py /workspace/shared/soc-analyst-copilot \
        --team "SOC-Copilot" --output ~/soc-copilot.docx

    # Generate .docx AND convert to PDF (requires LibreOffice)
    python3 scripts/code_to_docx.py /workspace/shared/soc-analyst-copilot \
        --team "SOC-Copilot" --output ~/soc-copilot.docx --pdf

    # Skip heavy directories
    python3 scripts/code_to_docx.py /workspace/shared/soc-analyst-copilot \
        --team "SOC-Copilot" --output ~/soc-copilot.docx --pdf \
        --ignore-dirs dist,node_modules,__pycache__,.git
"""

import argparse
import json
import os
import re
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# File extensions to include
CODE_EXTENSIONS = {
    ".py", ".ipynb", ".yaml", ".yml", ".json", ".sh", ".bash",
    ".txt", ".md", ".csv", ".toml", ".cfg", ".ini", ".conf",
    ".html", ".css", ".js", ".ts", ".jsx", ".tsx",
    ".c", ".cpp", ".h", ".hpp", ".java", ".go", ".rs",
    ".r", ".R", ".sql", ".dockerfile", ".env",
}

# Files/dirs to skip
SKIP_DIRS = {
    "__pycache__", ".git", ".ipynb_checkpoints", "node_modules",
    ".venv", "venv", ".tox", ".mypy_cache", ".pytest_cache", "dist",
}
SKIP_FILES = {".DS_Store", "Thumbs.db"}

# Max file size to include (500 KB)
MAX_FILE_SIZE = 500 * 1024


def should_include(filepath: str) -> bool:
    basename = os.path.basename(filepath)
    if basename in SKIP_FILES:
        return False
    _, ext = os.path.splitext(basename)
    if basename.lower() in ("dockerfile", "makefile", "requirements.txt"):
        return True
    return ext.lower() in CODE_EXTENSIONS


def collect_files(directory: str, extra_skip_dirs: set | None = None) -> list:
    skip = SKIP_DIRS | (extra_skip_dirs or set())
    files = []
    for root, dirs, filenames in os.walk(directory):
        dirs[:] = [d for d in dirs if d not in skip and not d.startswith(".")]
        dirs.sort()
        for fname in sorted(filenames):
            fpath = os.path.join(root, fname)
            if not should_include(fpath):
                continue
            try:
                if os.path.getsize(fpath) > MAX_FILE_SIZE:
                    continue
            except OSError:
                continue
            rel = os.path.relpath(fpath, directory)
            files.append((rel, fpath))
    return files


def _add_code_block(doc, code: str):
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn("w:rFonts"), {
            qn("w:ascii"): "Consolas",
            qn("w:hAnsi"): "Consolas",
            qn("w:cs"): "Courier New",
        })
        rPr.insert(0, rFonts)


def _read_notebook_cells(filepath: str) -> list:
    try:
        with open(filepath, encoding="utf-8", errors="replace") as f:
            nb = json.load(f)
    except (json.JSONDecodeError, ValueError) as e:
        return [{"cell_number": 1, "cell_type": "error", "source": f"[Could not parse notebook: {e}]"}]

    raw_cells = nb.get("cells", [])
    if not isinstance(raw_cells, list):
        return [{"cell_number": 1, "cell_type": "error", "source": "[Invalid notebook format]"}]

    cells = []
    for i, cell in enumerate(raw_cells, 1):
        cell_type = cell.get("cell_type", "code")
        source_raw = cell.get("source", [])
        source = "".join(source_raw) if isinstance(source_raw, list) else str(source_raw)
        output_texts = []
        for out in cell.get("outputs", []):
            if out.get("output_type") == "stream":
                output_texts.append("".join(out.get("text", [])))
            elif out.get("output_type") in ("execute_result", "display_data"):
                text_data = out.get("data", {}).get("text/plain", [])
                output_texts.append("".join(text_data) if isinstance(text_data, list) else text_data)
            elif out.get("output_type") == "error":
                tb = out.get("traceback", [])
                output_texts.append(re.sub(r"\x1b\[[0-9;]*m", "", "\n".join(tb)))
        cells.append({
            "cell_number": i, "cell_type": cell_type,
            "source": source, "outputs": output_texts,
        })
    return cells


def _add_notebook(doc, filepath: str):
    cells = _read_notebook_cells(filepath)
    for cell in cells:
        num, ctype, source = cell["cell_number"], cell["cell_type"], cell["source"]
        header_p = doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(8)
        header_p.paragraph_format.space_after = Pt(2)
        run = header_p.add_run(f"Cell {num}  [{ctype}]")
        run.font.size = Pt(9)
        run.font.bold = True
        if ctype == "markdown":
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
        elif ctype == "code":
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)

        if ctype == "markdown":
            for line in source.splitlines():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = Pt(14)
                if line.startswith("#"):
                    stripped = line.lstrip("#").strip()
                    level = min(len(line) - len(line.lstrip("#")), 4)
                    run = p.add_run(stripped)
                    run.font.bold = True
                    run.font.size = Pt(12 - level)
                else:
                    run = p.add_run(line if line else " ")
                    run.font.size = Pt(9)
        else:
            _add_code_block(doc, source)

        outputs = cell.get("outputs", [])
        if outputs:
            out_header = doc.add_paragraph()
            out_header.paragraph_format.space_before = Pt(4)
            out_header.paragraph_format.space_after = Pt(1)
            run = out_header.add_run("Output:")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
            combined = "\n".join(outputs)
            lines = combined.splitlines()
            if len(lines) > 50:
                combined = "\n".join(lines[:50]) + f"\n... ({len(lines) - 50} more lines truncated)"
            _add_code_block(doc, combined)


def generate_docx(directory: str, team_name: str, output_path: str, extra_skip_dirs: set | None = None):
    files = collect_files(directory, extra_skip_dirs)
    if not files:
        print(f"No code files found in {directory}")
        sys.exit(1)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    # Title page
    doc.add_paragraph()
    title = doc.add_heading("Code Submission", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(team_name)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xED, 0x1C, 0x24)  # AMD red

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("TCS & AMD AI Hackathon 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(f"\n{len(files)} file(s) included")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Table of contents
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for i, (rel_path, _) in enumerate(files, 1):
        p = doc.add_paragraph(f"{i}. {rel_path}", style="List Number")
        p.paragraph_format.space_after = Pt(2)

    # Each file
    for i, (rel_path, abs_path) in enumerate(files, 1):
        doc.add_page_break()
        doc.add_heading(f"{rel_path}", level=2)

        size = os.path.getsize(abs_path)
        size_str = f"{size:,} bytes" if size < 1024 else f"{size / 1024:.1f} KB"
        meta = doc.add_paragraph()
        run = meta.add_run(f"File {i} of {len(files)}  |  {size_str}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

        doc.add_paragraph("_" * 80).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        try:
            if rel_path.endswith(".ipynb"):
                _add_notebook(doc, abs_path)
            else:
                with open(abs_path, encoding="utf-8", errors="replace") as f:
                    content = f.read()
                _add_code_block(doc, content)
        except Exception as e:
            _add_code_block(doc, f"[Error reading file: {e}]")

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"[✓] {output_path}  ({len(files)} files, {size_kb:.0f} KB)")
    return output_path


def convert_to_pdf(docx_path: str) -> str:
    """Convert a .docx to .pdf using LibreOffice headless."""
    pdf_path = docx_path.replace(".docx", ".pdf")
    outdir   = os.path.dirname(os.path.abspath(docx_path)) or "."

    # Try LibreOffice (standard on Ubuntu)
    for cmd in ("libreoffice", "libreoffice7.6", "libreoffice7.5", "soffice"):
        if _which(cmd):
            print(f"[•] Converting to PDF via {cmd}…")
            result = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf",
                 os.path.abspath(docx_path), "--outdir", outdir],
                capture_output=True, text=True,
            )
            if result.returncode == 0:
                print(f"[✓] {pdf_path}")
                return pdf_path
            else:
                print(f"[!] LibreOffice error: {result.stderr.strip()}", file=sys.stderr)
                break

    # Try python docx2pdf as fallback
    try:
        from docx2pdf import convert
        print("[•] Converting to PDF via docx2pdf…")
        convert(docx_path, pdf_path)
        print(f"[✓] {pdf_path}")
        return pdf_path
    except ImportError:
        pass

    print(
        "\n[!] Could not convert to PDF automatically.\n"
        "    Install LibreOffice:  apt-get install -y libreoffice\n"
        "    Then re-run with --pdf, or convert manually:\n"
        f"    libreoffice --headless --convert-to pdf {docx_path}",
        file=sys.stderr,
    )
    return docx_path


def _which(cmd: str) -> bool:
    import shutil
    return shutil.which(cmd) is not None


def main():
    parser = argparse.ArgumentParser(
        description="Generate a Word doc (+ optional PDF) from a code directory.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("directory", help="Root directory containing code files")
    parser.add_argument("--team",        default="Team",  help="Team name on the title page")
    parser.add_argument("--output", "-o", default=None,   help="Output .docx path (default: <team>.docx)")
    parser.add_argument("--pdf",  action="store_true",    help="Also convert the .docx to PDF via LibreOffice")
    parser.add_argument("--ignore-dirs", default=None,    help="Extra dirs to skip, comma-separated (e.g. data,weights)")
    args = parser.parse_args()

    if not os.path.isdir(args.directory):
        print(f"Error: {args.directory!r} is not a directory")
        sys.exit(1)

    extra_skip = None
    if args.ignore_dirs:
        extra_skip = {d.strip() for d in args.ignore_dirs.split(",") if d.strip()}

    output = args.output or f"{args.team.replace(' ', '_')}.docx"
    if not output.endswith(".docx"):
        output += ".docx"

    docx_path = generate_docx(args.directory, args.team, output, extra_skip)

    if args.pdf:
        convert_to_pdf(docx_path)


if __name__ == "__main__":
    main()
